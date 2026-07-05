"""Unit tests for the document parser."""

import pytest

from app.modules.rag.infrastructure.parser.document_parser import (
    UnsupportedFormatError,
    parse_to_markdown,
)


def test_parses_markdown_and_text() -> None:
    assert parse_to_markdown("notes.md", b"# Title\n\nbody") == "# Title\n\nbody"
    assert parse_to_markdown("plain.txt", b"just text") == "just text"


def test_unknown_extension_raises() -> None:
    with pytest.raises(UnsupportedFormatError):
        parse_to_markdown("archive.zip", b"...")


def test_docx_converts_headings_lists_and_tables_to_markdown() -> None:
    from io import BytesIO

    from docx import Document

    document = Document()
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("An intro paragraph.")
    document.add_paragraph("First point", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "Name", "Score"
    table.cell(1, 0).text, table.cell(1, 1).text = "Ana", "10"
    buffer = BytesIO()
    document.save(buffer)

    markdown = parse_to_markdown("doc.docx", buffer.getvalue())

    assert "# Chapter 1" in markdown
    assert "- First point" in markdown
    assert "| Name | Score |" in markdown
    assert "| Ana | 10 |" in markdown


def test_xlsx_converts_sheet_to_markdown_table() -> None:
    from io import BytesIO

    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Grades"
    sheet.append(["Name", "Score"])
    sheet.append(["Ana", 10])
    buffer = BytesIO()
    workbook.save(buffer)

    markdown = parse_to_markdown("sheet.xlsx", buffer.getvalue())

    assert "## Grades" in markdown
    assert "| Name | Score |" in markdown
    assert "| Ana | 10 |" in markdown
